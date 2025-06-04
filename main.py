from tkinter import ttk
import tkinter as tk
import tkinter.messagebox as tkmb
import openpyxl
from openpyxl import *
from datetime import datetime
from tkinter import filedialog
import ttkbootstrap as ttkb
from ttkbootstrap.constants import *
from openpyxl.utils import range_boundaries
import docx
from docx.shared import Inches
import math
wb = openpyxl.load_workbook('Sizing_template.xlsx')
sheet = wb.active
wb2 = openpyxl.load_workbook('Standard costing sheet.xlsx', data_only=True)
sheet2 = wb2.active

root=ttkb.Window(themename="darkly")
root.title("Sizing System")
root.geometry("960x860")
root.resizable(False,False)

global ups_make_entry, ups_model_entry, ups_rating_entry, actual_loadkva_entry, actual_loadkw_entry
global power_factor_entry, inverter_efficiency_entry, nominal_dc_voltage_entry, backup_requirement_entry
global cell_chemisrty_combobox

def mainscreen():
    title_label = ttkb.Label(root, text="Sizing Software", font=("Segoe UI", 32, "bold"))
    title_label.place(relx=0.5, rely=0.1, anchor="center")
    def next():
        global customername
        customername=customer_entry.get()
        global providername
        providername=provider_entry.get()
        global date
        date=date_entry.get()
        customer_name = customer_entry.get()
        solution_provider = provider_entry.get()
        date = date_entry.get()
        if not customer_name or not solution_provider or not date:
            tkmb.showerror("Error", "Please fill in all fields.")
            return
        try:
            datetime.strptime(date, "%d-%m-%Y")
            input()
        except ValueError:
            tkmb.showerror("Error", "Date must be in DD-MM-YYYY format.")
            return
    global center_frame
    center_frame = ttkb.Frame(root)
    center_frame.place(relx=0.5, rely=0.5, anchor="center")
    customer_label = ttkb.Label(center_frame, text="Customer Name:", font=("Segoe UI", 14))
    customer_label.grid(row=0, column=0, padx=10, pady=10, sticky="e")
    customer_entry = ttkb.Entry(center_frame, font=("Segoe UI", 14), width=30)
    customer_entry.grid(row=0, column=1, padx=10, pady=10)
    provider_label = ttkb.Label(center_frame, text="Solution Provider:", font=("Segoe UI", 14))
    provider_label.grid(row=1, column=0, padx=10, pady=10, sticky="e")
    provider_entry = ttkb.Entry(center_frame, font=("Segoe UI", 14), width=30)
    provider_entry.grid(row=1, column=1, padx=10, pady=10)
    date_label = ttkb.Label(center_frame, text="Date:", font=("Segoe UI", 14))
    date_label.grid(row=2, column=0, padx=10, pady=10, sticky="e")
    date_entry = ttkb.Entry(center_frame, font=("Segoe UI", 14), width=30)
    date_entry.insert(0, datetime.now().strftime("%d-%m-%Y"))
    date_entry.grid(row=2, column=1, padx=10, pady=10)
    next_button = ttkb.Button(center_frame, text="Next", command=lambda: next())
    next_button.grid(row=3, column=0, columnspan=2, pady=20)
    search_button = ttkb.Button(center_frame, text="Search", command=costing_screen)
    search_button.grid(row=4, column=0, columnspan=2, pady=10)

def input():
    try:
        costing_frame.destroy()
    except:
        pass
    global ups_make_entry, ups_model_entry, ups_rating_entry, actual_loadkva_entry, actual_loadkw_entry
    global power_factor_entry, inverter_efficiency_entry, nominal_dc_voltage_entry, backup_requirement_entry
    global cell_chemisrty_combobox
    global input_frame
    def size():
        global ups_make
        global ups_model
        global ups_rating
        global actual_loadkva
        global power_factor
        global inverter_efficiency
        global nominal_dc_voltage
        global backup_requirement
        global calc_load
        global actual_loadkw
        global maximum_charging_voltage
        global endcell_voltage
        global energy_required
        global capacity_required
        global noofcells
        global cell_chemisrty
        ups_make = ups_make_entry.get()
        ups_model = ups_model_entry.get()
        ups_rating = float(ups_rating_entry.get())
        actual_loadkva = float(actual_loadkva_entry.get())
        actual_loadkw = float(actual_loadkw_entry.get())
        power_factor = float(power_factor_entry.get())
        inverter_efficiency = float(inverter_efficiency_entry.get())
        nominal_dc_voltage = float(nominal_dc_voltage_entry.get())
        backup_requirement = float(backup_requirement_entry.get())
        cell_chemisrty = cell_chemisrty_combobox.get()
        if cell_chemisrty == "LFP":
            nominal_vdc = 3.2
        elif cell_chemisrty == "NPM":
            nominal_vdc = 3.6
        if actual_loadkw:
            calc_load = round((actual_loadkw/inverter_efficiency),1)
        elif actual_loadkva:
            calc_load = round(((actual_loadkva * power_factor)/inverter_efficiency),1)
        elif ups_rating:
            calc_load = round(((ups_rating * power_factor)/inverter_efficiency),1)
        if nominal_dc_voltage == 12:
            noofcells = 4
        elif nominal_dc_voltage == 24:
            noofcells = 8
        elif nominal_dc_voltage == 36:
            noofcells = 11
        elif nominal_dc_voltage == 48:
            noofcells = 15
        elif nominal_dc_voltage == 72:
            noofcells = 23
        elif nominal_dc_voltage == 96:
            noofcells = 30
        elif nominal_dc_voltage == 120:
            noofcells = 38
        elif nominal_dc_voltage == 144:
            noofcells = 45
        elif nominal_dc_voltage == 192:
            noofcells = 60
        elif nominal_dc_voltage == 240:
            noofcells = 75
        elif nominal_dc_voltage == 336:
            noofcells = 105
        elif nominal_dc_voltage == 360:
            noofcells = 112
        elif nominal_dc_voltage == 384:
            noofcells = 120
        elif nominal_dc_voltage == 408:
            noofcells = 128
        elif nominal_dc_voltage == 480:
            noofcells = 150
        elif nominal_dc_voltage == 512:
            noofcells = 160
        elif nominal_dc_voltage == 528:
            noofcells = 165
        elif nominal_dc_voltage == 576:
            noofcells = 180
        global maximum_charging_voltage_calc
        global endcell_voltage_calc
        global energy_required_calc
        global capacity_required_calc
        maximum_charging_voltage_calc = (noofcells* 3.6)
        endcell_voltage_calc = noofcells* 2.8
        energy_required_calc = (calc_load * backup_requirement) / 60
        capacity_required_calc = (energy_required_calc * 1000) / endcell_voltage_calc
        maximum_charging_voltage = maximum_charging_voltage_calc
        endcell_voltage = round(endcell_voltage_calc,1)
        energy_required = round(energy_required_calc,1)
        capacity_required = round(capacity_required_calc,1)
        input2()
    center_frame.destroy()
    input_frame = ttkb.Frame(root)
    input_frame.place(relx=0.5, rely=0.5, anchor="center")
    ups_make_label = ttkb.Label(input_frame, text="UPS Make:", font=("Segoe UI", 12))
    ups_make_label.grid(row=0, column=0, padx=10, pady=5, sticky="e")
    ups_make_entry = ttkb.Entry(input_frame, font=("Segoe UI", 12), width=30)
    ups_make_entry.grid(row=0, column=1, padx=10, pady=5)
    ups_model_label = ttkb.Label(input_frame, text="UPS Model:", font=("Segoe UI", 12))
    ups_model_label.grid(row=1, column=0, padx=10, pady=5, sticky="e")
    ups_model_entry = ttkb.Entry(input_frame, font=("Segoe UI", 12), width=30)
    ups_model_entry.grid(row=1, column=1, padx=10, pady=5)
    ups_rating_label = ttkb.Label(input_frame, text="UPS Rating (kVA):", font=("Segoe UI", 12))
    ups_rating_label.grid(row=2, column=0, padx=10, pady=5, sticky="e")
    ups_rating_entry = ttkb.Entry(input_frame, font=("Segoe UI", 12), width=30)
    ups_rating_entry.insert(0, "0")
    ups_rating_entry.grid(row=2, column=1, padx=10, pady=5)
    actual_loadkva_label = ttkb.Label(input_frame, text="Actual Load (kVA):", font=("Segoe UI", 12))
    actual_loadkva_label.grid(row=3, column=0, padx=10, pady=5, sticky="e")
    actual_loadkva_entry = ttkb.Entry(input_frame, font=("Segoe UI", 12), width=30)
    actual_loadkva_entry.insert(0, "0")
    actual_loadkva_entry.grid(row=3, column=1, padx=10, pady=5)
    actual_loadkw_label = ttkb.Label(input_frame, text="Actual Load (kW):", font=("Segoe UI", 12))
    actual_loadkw_label.grid(row=4, column=0, padx=10, pady=5, sticky="e")
    actual_loadkw_entry = ttkb.Entry(input_frame, font=("Segoe UI", 12), width=30)
    actual_loadkw_entry.insert(0, "0")
    actual_loadkw_entry.grid(row=4, column=1, padx=10, pady=5)
    power_factor_label = ttkb.Label(input_frame, text="Power Factor:", font=("Segoe UI", 12))
    power_factor_label.grid(row=5, column=0, padx=10, pady=5, sticky="e")
    power_factor_entry = ttkb.Entry(input_frame, font=("Segoe UI", 12), width=30)
    power_factor_entry.insert(0, "0")
    power_factor_entry.grid(row=5, column=1, padx=10, pady=5)
    inverter_efficiency_label = ttkb.Label(input_frame, text="Inverter Efficiency:", font=("Segoe UI", 12))
    inverter_efficiency_label.grid(row=6, column=0, padx=10, pady=5, sticky="e")
    inverter_efficiency_entry = ttkb.Entry(input_frame, font=("Segoe UI", 12), width=30)
    inverter_efficiency_entry.insert(0, "0")
    inverter_efficiency_entry.grid(row=6, column=1, padx=10, pady=5)
    nominal_dc_voltage_label = ttkb.Label(input_frame, text="Nominal DC Voltage (V):", font=("Segoe UI", 12))
    nominal_dc_voltage_label.grid(row=7, column=0, padx=10, pady=5, sticky="e")
    nominal_dc_voltage_entry = ttkb.Entry(input_frame, font=("Segoe UI", 12), width=30)
    nominal_dc_voltage_entry.insert(0, "0")
    nominal_dc_voltage_entry.grid(row=7, column=1, padx=10, pady=5)
    backup_requirement_label = ttkb.Label(input_frame, text="Backup Requirement (Min):", font=("Segoe UI", 12))
    backup_requirement_label.grid(row=8, column=0, padx=10, pady=5, sticky="e")
    backup_requirement_entry = ttkb.Entry(input_frame, font=("Segoe UI", 12), width=30)
    backup_requirement_entry.insert(0, "0")
    backup_requirement_entry.grid(row=8, column=1, padx=10, pady=5)
    cell_chemisrty_label = ttkb.Label(input_frame, text="Cell Chemistry:", font=("Segoe UI", 12))
    cell_chemisrty_label.grid(row=9, column=0, padx=10, pady=5, sticky="e")
    cell_chemisrty_var = tk.StringVar()
    cell_chemisrty_combobox = ttkb.Combobox(input_frame, textvariable=cell_chemisrty_var, font=("Segoe UI", 12), width=28, state="readonly")
    cell_chemisrty_combobox['values'] = ("LFP", "NPM")
    cell_chemisrty_combobox.current(0)
    cell_chemisrty_combobox.grid(row=9, column=1, padx=10, pady=5)
    size_button = ttkb.Button(input_frame, text="Size", width=20,command=lambda: size())
    size_button.grid(row=10, column=0, columnspan=2, pady=20)

def input2():
    global input_frame2
    global offered_battery_config_entry
    global calc_load_entry
    global noofcells_entry
    global maximum_charging_voltage_entry
    global endcell_voltage_entry
    global minimum_charging_current_entry
    global maximum_charging_current_entry
    global energy_required_entry
    global capacity_required_entry
    global nearest_available_capacity_entry
    global total_available_energy_entry
    global backup_time_entry

    input_frame.destroy()
    input_frame2 = ttkb.Frame(root)
    input_frame2.place(relx=0.5, rely=0.5, anchor="center")
    calc_load_label = ttkb.Label(input_frame2, text="Calculated Load in kW:", font=("Segoe UI", 12))
    calc_load_label.grid(row=0, column=0, padx=10, pady=5, sticky="e")
    calc_load_entry = ttkb.Entry(input_frame2, font=("Segoe UI", 12), width=30)
    calc_load_entry.insert(0,calc_load)
    calc_load_entry.grid(row=0, column=1, padx=10, pady=5)
    noofcells_label = ttkb.Label(input_frame2, text="Number of Cells:", font=("Segoe UI", 12))
    noofcells_label.grid(row=2, column=0, padx=10, pady=5, sticky="e")
    noofcells_entry = ttkb.Entry(input_frame2, font=("Segoe UI", 12), width=30)
    noofcells_entry.insert(0,noofcells)
    noofcells_entry.grid(row=2, column=1, padx=10, pady=5)
    maximum_charging_voltage_label = ttkb.Label(input_frame2, text="Maximum Charging Voltage(V):", font=("Segoe UI", 12))
    maximum_charging_voltage_label.grid(row=3, column=0, padx=10, pady=5, sticky="e")
    maximum_charging_voltage_entry = ttkb.Entry(input_frame2, font=("Segoe UI", 12), width=30)
    maximum_charging_voltage_entry.insert(0,maximum_charging_voltage)
    maximum_charging_voltage_entry.grid(row=3, column=1, padx=10, pady=5)
    endcell_voltage_label = ttkb.Label(input_frame2, text="End Cell Voltage(V):", font=("Segoe UI", 12))
    endcell_voltage_label.grid(row=4, column=0, padx=10, pady=5, sticky="e")
    endcell_voltage_entry = ttkb.Entry(input_frame2, font=("Segoe UI", 12), width=30)
    endcell_voltage_entry.insert(0,endcell_voltage)
    endcell_voltage_entry.grid(row=4, column=1, padx=10, pady=5)
    minimum_charging_current_label = ttkb.Label(input_frame2, text="Minimum Charging Current:", font=("Segoe UI", 12))
    minimum_charging_current_label.grid(row=5, column=0, padx=10, pady=5, sticky="e")
    minimum_charging_current_entry = ttkb.Entry(input_frame2, font=("Segoe UI", 12), width=30)
    minimum_charging_current_entry.grid(row=5, column=1, padx=10, pady=5)
    minimum_charging_current_entry.insert(0,"0.1C")
    maximum_charging_current_label = ttkb.Label(input_frame2, text="Maximum Charging Current:", font=("Segoe UI", 12))
    maximum_charging_current_label.grid(row=6, column=0, padx=10, pady=5, sticky="e")
    maximum_charging_current_entry = ttkb.Entry(input_frame2, font=("Segoe UI", 12), width=30)
    maximum_charging_current_entry.grid(row=6, column=1, padx=10, pady=5)
    maximum_charging_current_entry.insert(0,"0.5C")
    energy_required_label = ttkb.Label(input_frame2, text="Energy Required(kWh):", font=("Segoe UI", 12))
    energy_required_label.grid(row=7, column=0, padx=10, pady=5, sticky="e")
    energy_required_entry = ttkb.Entry(input_frame2, font=("Segoe UI", 12), width=30)
    energy_required_entry.insert(0,energy_required)
    energy_required_entry.grid(row=7, column=1, padx=10, pady=5)
    ageing_label = ttkb.Label(input_frame2, text="Ageing:", font=("Segoe UI", 12))
    ageing_label.grid(row=8, column=0, padx=10, pady=5, sticky="e")
    ageing_entry = ttkb.Entry(input_frame2, font=("Segoe UI", 12), width=30)
    ageing_entry.insert(0, 0)
    ageing_entry.grid(row=8, column=1, padx=10, pady=5)
    def on_ageing_change(event):
        age_str = ageing_entry.get()
        try:
            age = float(age_str)
        except ValueError:
            capacity_required_entry.delete(0, tk.END)
            return
        capacity_required_calc = ((energy_required_calc * 1000) / endcell_voltage_calc) + ((energy_required_calc * 1000) / endcell_voltage_calc) * age
        capacity_required = round(capacity_required_calc, 1)
        capacity_required_entry.delete(0, tk.END)
        capacity_required_entry.insert(0, capacity_required)
        try:
            backup_time = (backup_requirement / capacity_required_calc) * float(nearest_available_capacity_entry.get())
            backup_time_entry.delete(0, tk.END)
            backup_time_entry.insert(0, round(backup_time,1))
        except ValueError:
            backup_time_entry.delete(0, tk.END)
            return
    ageing_entry.bind("<KeyRelease>", on_ageing_change)
    capacity_required_label = ttkb.Label(input_frame2, text="Capacity Required(Ah):", font=("Segoe UI", 12))
    capacity_required_label.grid(row=9, column=0, padx=10, pady=5, sticky="e")
    capacity_required_entry = ttkb.Entry(input_frame2, font=("Segoe UI", 12), width=30)
    capacity_required_entry.insert(0,capacity_required)
    capacity_required_entry.grid(row=9, column=1, padx=10, pady=5)
    nearest_available_capacity_label = ttkb.Label(input_frame2, text="Nearest Available Capacity(Ah):", font=("Segoe UI", 12))
    nearest_available_capacity_label.grid(row=10, column=0, padx=10, pady=5, sticky="e")
    nearest_available_capacity_entry = ttkb.Entry(input_frame2, font=("Segoe UI", 12), width=30)
    nearest_available_capacity_entry.grid(row=10, column=1, padx=10, pady=5)
    def on_nearest_capacity_change(event):
        global value_str
        global total_available_energy
        global backup_time
        value_str = nearest_available_capacity_entry.get()
        try:
            value = float(value_str)
        except ValueError:
            backup_time_entry.delete(0, tk.END)
            total_available_energy_entry.delete(0, tk.END)
            return
        backup_time = (backup_requirement / float(capacity_required_entry.get())) * value
        backup_time_entry.delete(0, tk.END)
        backup_time_entry.insert(0, math.floor(backup_time))
        total_available_energy = ((nominal_dc_voltage * value) / 1000)
        total_available_energy_entry.delete(0, tk.END)
        total_available_energy_entry.insert(0,total_available_energy)
        offered_battery_config=f"{int(nominal_dc_voltage)}V {int(value)}Ah"
        offered_battery_config_entry.delete(0, tk.END)
        offered_battery_config_entry.insert(0, offered_battery_config)
    nearest_available_capacity_entry.bind("<KeyRelease>", on_nearest_capacity_change)
    offered_battery_config_label = ttkb.Label(input_frame2, text="Offered Battery Configuration:", font=("Segoe UI", 12))
    offered_battery_config_label.grid(row=11, column=0, padx=10, pady=5, sticky="e")
    offered_battery_config_entry = ttkb.Entry(input_frame2, font=("Segoe UI", 12), width=30)
    offered_battery_config_entry.grid(row=11, column=1, padx=10, pady=5)
    total_available_energy_label = ttkb.Label(input_frame2, text="Total Available Energy(kWh):", font=("Segoe UI", 12))
    total_available_energy_label.grid(row=12, column=0, padx=10, pady=5, sticky="e")
    total_available_energy_entry = ttkb.Entry(input_frame2, font=("Segoe UI", 12), width=30)
    total_available_energy_entry.grid(row=12, column=1, padx=10, pady=5)
    backup_time_label = ttkb.Label(input_frame2, text="Backup Time(Aproximate Minutes):", font=("Segoe UI", 12))
    backup_time_label.grid(row=13, column=0, padx=10, pady=5, sticky="e")
    backup_time_entry = ttkb.Entry(input_frame2, font=("Segoe UI", 12), width=30)
    backup_time_entry.grid(row=13, column=1, padx=10, pady=5)
    def save_to_excel():
        if actual_loadkva:
            actual_load= actual_loadkva
        elif actual_loadkw:
            actual_load= actual_loadkw
        elif ups_rating:
            actual_load= 0
        data = {
            "Customer Name": customername,
            "Solution Provider": providername,
            "Date": date,
            "UPS Make": ups_make,
            "UPS Model": ups_model,
            "UPS Rating (KVA)": ups_rating,
            "Actual Load (KVA)": actual_load,
            "Power Factor": power_factor,
            "Inverter Efficiency": inverter_efficiency,
            "Nominal DC Voltage (V)": nominal_dc_voltage,
            "Backup Requirement (Min)": backup_requirement,
            "Cell Chemistry": cell_chemisrty,
            "Calculated Load (kW)": calc_load_entry.get(),
            "Maximum Charging Voltage(V)": maximum_charging_voltage_entry.get(),
            "End Cell Voltage(V)": endcell_voltage_entry.get(),
            "Minimum Charging Current": minimum_charging_current_entry.get(),
            "Maximum Charging Current": maximum_charging_current_entry.get(),
            "Energy Required(kWh)": energy_required_entry.get(),
            "Capacity Required(Ah)": capacity_required_entry.get(),
            "Nearest Available Capacity(Ah)": nearest_available_capacity_entry.get(),
            "Offered Battery Configuration": offered_battery_config_entry.get(),
            "Total Available Energy(kWh)": total_available_energy_entry.get(),
            "Backup Time(Aproximate minimum)": backup_time_entry.get()
        }   
        cell_map = {
            "Customer Name": "C4",
            "Solution Provider": "C5",
            "Date": "H3",
            "UPS Make": "A8",
            "UPS Model": "B6",
            "UPS Rating (KVA)": "C8",
            "Actual Load (KVA)": "D8",
            "Power Factor": "E8",
            "Inverter Efficiency": "F8",
            "Nominal DC Voltage (V)": "G8",
            "Backup Requirement (Min)": "H8",
            "Cell Chemistry": "E11",
            "Calculated Load (kW)": "E12",
            "Maximum Charging Voltage(V)": "E13",
            "End Cell Voltage(V)": "E14",
            "Minimum Charging Current": "E15",
            "Maximum Charging Current": "E16",
            "Energy Required(kWh)": "E17",
            "Capacity Required(Ah)": "E18",
            "Nearest Available Capacity(Ah)": "E19",
            "Offered Battery Configuration": "E20",
            "Total Available Energy(kWh)": "E21",
            "Backup Time(Aproximate minimum)": "E22"
        }
        merged_ranges = list(sheet.merged_cells.ranges)
        for key, cell in cell_map.items():
            merged = None
            for mrange in merged_ranges:
                if cell in mrange:
                    merged = mrange
                    break
            if merged:
                sheet.unmerge_cells(str(merged))
                sheet[cell] = data[key]
                sheet.merge_cells(str(merged))
            else:
                sheet[cell] = data[key]
        save_path = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel files", "*.xlsx")])
        if save_path:
            wb.save(save_path)
            tkmb.showinfo("Success", "Data saved to Excel successfully.")
    save_button = ttkb.Button(input_frame2, text="Save to Excel", command=save_to_excel)
    save_button.grid(row=14, column=0, pady=20, sticky="e")
    clear_button = ttkb.Button(input_frame2, text="Make Costing", command=costing_screen)
    clear_button.grid(row=14, column=1, pady=20, sticky="w")

    def back_to_input():
        input_frame2.destroy()
        input()
        ups_make_entry.insert(0, ups_make)
        ups_model_entry.insert(0, ups_model)
        ups_rating_entry.insert(0, ups_rating)
        actual_loadkva_entry.insert(0, actual_loadkva)
        actual_loadkw_entry.insert(0, actual_loadkw)
        power_factor_entry.insert(0, power_factor)
        inverter_efficiency_entry.insert(0, inverter_efficiency)
        nominal_dc_voltage_entry.insert(0, nominal_dc_voltage)
        backup_requirement_entry.insert(0, backup_requirement)       
        if cell_chemisrty == "LFP":
            cell_chemisrty_combobox.current(0)
        else:
            cell_chemisrty_combobox.current(1)
        cell_chemisrty_combobox.grid(row=9, column=1, padx=10, pady=5)

    back_button = ttkb.Button(input_frame2, text="Back", command=back_to_input)
    back_button.grid(row=14, column=2, pady=20, sticky="w")

def costing_screen():
    global costing_frame
    def export_costing():
        template_wb = openpyxl.load_workbook('costing_sheet_template.xlsx')
        template_ws = template_wb.active
        option_count = 0
        for col in range(1, 4):
            if any(tree.set(item, f"Option {col}") for item in tree.get_children()):
                option_count += 1
        for col in range(2, 2 + option_count):
            template_ws.cell(row=3, column=col).value = backup_time_var.get()
            template_ws.cell(row=4, column=col).value = battery_config_entry.get()
        for idx, item_id in enumerate(tree.get_children()):
            excel_row = idx + 5
            if excel_row > 45:
                break
            values = tree.item(item_id, 'values')
            for col_offset, val in enumerate(values[1:1+option_count], start=2):
                template_ws.cell(row=excel_row, column=col_offset).value = val
        save_path = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel files", "*.xlsx")])
        if save_path:
            template_wb.save(save_path)
            tkmb.showinfo("Success", "Costing exported successfully.")
    global offered_battery_config_value
    try:
        offered_battery_config_value = offered_battery_config_entry.get()
        input_frame2.destroy()
    except:
        offered_battery_config_value = ""
    try:
        center_frame.destroy()
    except:
        pass
    costing_frame = ttkb.Frame(root)
    costing_frame.place(relx=0.5, rely=0.5, anchor="center")
    backup_time_label = ttkb.Label(costing_frame, text="Select Backup Time:", font=("Segoe UI", 12))
    backup_time_label.grid(row=1, column=0, padx=10, pady=10, sticky="e")
    backup_time_var = tk.StringVar(value="")  # Set default to empty
    battery_config_label = ttkb.Label(costing_frame, text="Battery Configuration:", font=("Segoe UI", 12))
    battery_config_label.grid(row=0, column=0, padx=10, pady=10, sticky="e")
    battery_config_entry = ttkb.Entry(costing_frame, font=("Segoe UI", 12), width=30)
    battery_config_entry.insert(0, offered_battery_config_value)
    battery_config_entry.grid(row=0, column=1, padx=10, pady=10)
    backup_time_combobox = ttkb.Combobox(costing_frame, textvariable=backup_time_var, font=("Segoe UI", 12), width=28, state="readonly")
    backup_time_combobox['values'] = ("15min", "30min", "60min", "120min")
    # Do NOT call backup_time_combobox.current(0)
    backup_time_combobox.grid(row=1, column=1, padx=10, pady=10)
    def on_backup_time_select(event):
        global sheet3
        selection = backup_time_var.get()
        sheet3 = wb2[selection]
        found_columns = []
        search_value = battery_config_entry.get().strip().lower()
        for col in range(1, sheet3.max_column + 1):
            cell_value = sheet3.cell(row=4, column=col).value
            if cell_value and str(cell_value).strip().lower() == search_value:
                col_letter = openpyxl.utils.get_column_letter(col)
                found_columns.append(col_letter)
        for idx in range(len(tree.get_children())):
            tree.set(tree.get_children()[idx], "Option 1", "")
            tree.set(tree.get_children()[idx], "Option 2", "")
            tree.set(tree.get_children()[idx], "Option 3", "")
        for option_idx, col_letter in enumerate(found_columns[:3]):
            for row_idx, excel_row in enumerate(range(5, 46)):
                value = sheet3[f"{col_letter}{excel_row}"].value
                if isinstance(value, (float, int)):
                    value = round(value, 2)
                tree.set(tree.get_children()[row_idx], f"Option {option_idx+1}", value if value is not None else "")
        highlight_excel_rows = [21, 34, 17, 40, 42]
        global prices, centretapping
        prices = []
        for option_idx, col_letter in enumerate(found_columns[:3]):
            price_cell_value = sheet3[f"{col_letter}40"].value
            prices.append(price_cell_value)
        centretapping=[]
        for option_idx, col_letter in enumerate(found_columns[:3]):
            centretapping_cell_value = sheet3[f"{col_letter}1"].value
            centretapping.append(centretapping_cell_value)
        highlight_tree_indices = [row - row_start for row in highlight_excel_rows]
        tree.tag_configure('highlight', background='yellow', foreground='black')
        for idx in highlight_tree_indices:
            item_id = tree.get_children()[idx]
            tree.item(item_id, tags=('highlight',))
    backup_time_combobox.bind("<<ComboboxSelected>>", on_backup_time_select)
    row_start = 5
    row_end = 45
    row_labels = []
    for i in range(row_start, row_end + 1):
        cell_value = sheet2[f"A{i}"].value
        row_labels.append(cell_value if cell_value is not None else "")
    columns = ("Description", "Option 1", "Option 2", "Option 3")
    global tree
    tree = ttkb.Treeview(costing_frame, columns=columns, show="headings", height=15)
    # Set columns as wide as possible
    tree.heading("Description", text="Description")
    tree.column("Description", width=300, anchor="w", stretch=True)
    for col in columns[1:]:
        tree.heading(col, text=col)
        tree.column(col, width=200, anchor="center", stretch=True)
    for label in row_labels:
        tree.insert("", "end", values=(label, "", "", ""))
    vsb = ttkb.Scrollbar(costing_frame, orient="vertical", command=tree.yview)
    tree.configure(yscrollcommand=vsb.set)
    vsb.grid(row=2, column=4, sticky="ns")
    hsb = ttkb.Scrollbar(costing_frame, orient="horizontal", command=tree.xview)
    tree.configure(xscrollcommand=hsb.set)
    hsb.grid(row=3, column=0, columnspan=3, sticky="ew")
    tree.grid(row=2, column=0, columnspan=3, padx=10, pady=10, sticky="nsew")
    costing_frame.grid_rowconfigure(1, weight=1)
    costing_frame.grid_columnconfigure(0, weight=1)
    costing_frame.grid_columnconfigure(1, weight=1)
    button_frame3= ttkb.Frame(costing_frame)
    button_frame3.grid(row=5, column=0, pady=20)
    new_costing_button = ttkb.Button(button_frame3, text="New Costing", command=new_costing)
    new_costing_button.grid(row=0, column=0, padx=20)
    export_button = ttkb.Button(button_frame3, text="Export Costing", command=export_costing,bootstyle="success")
    export_button.grid(row=0, column=1, padx=20)
    def back_to_input2():
        costing_frame.destroy()
        input2()
        calc_load_entry.delete(0, tk.END)
        noofcells_entry.delete(0, tk.END)
        maximum_charging_voltage_entry.delete(0, tk.END)
        endcell_voltage_entry.delete(0, tk.END)
        minimum_charging_current_entry.delete(0, tk.END)
        maximum_charging_current_entry.delete(0, tk.END)
        energy_required_entry.delete(0, tk.END)
        capacity_required_entry.delete(0, tk.END)
        nearest_available_capacity_entry.delete(0, tk.END)
        total_available_energy_entry.delete(0, tk.END)
        backup_time_entry.delete(0, tk.END)
        calc_load_entry.insert(0, calc_load)
        noofcells_entry.insert(0, noofcells)
        maximum_charging_voltage_entry.insert(0, maximum_charging_voltage)
        endcell_voltage_entry.insert(0, endcell_voltage)
        minimum_charging_current_entry.insert(0, "0.1C")
        maximum_charging_current_entry.insert(0, "0.5C")
        energy_required_entry.insert(0, energy_required)
        capacity_required_entry.insert(0, capacity_required)
        nearest_available_capacity_entry.insert(0,value_str)
        offered_battery_config_entry.insert(0, offered_battery_config_value)
        total_available_energy_entry.insert(0, total_available_energy)
        backup_time_entry.insert(0, backup_time)
    button_frame = ttkb.Frame(costing_frame)
    button_frame.grid(row=4, column=1 ,pady=10,columnspan=2)
    option1_button = ttkb.Button(button_frame, text="Add To Quotation", command=option1,bootstyle="warning")
    option1_button.grid(row=0, column=0, padx=40)
    option2_button = ttkb.Button(button_frame, text="Add To Quotation", command=option2,bootstyle="warning")
    option2_button.grid(row=0, column=1, padx=40)
    option3_button = ttkb.Button(button_frame, text="Add To Quotation", command=option3,bootstyle="warning")
    option3_button.grid(row=0, column=2, padx=40)
    make_quote_button = ttkb.Button(costing_frame, text="Make Quote", command=quotation,bootstyle="success")
    make_quote_button.grid(row=5, column=1, pady=20)
    button_frame2= ttkb.Frame(costing_frame)
    button_frame2.grid(row=5, column=2, pady=20)
    back_button = ttkb.Button(button_frame2, text="Back", command=back_to_input2,bootstyle="danger")
    back_button.grid(row=0, column=0)
    new_sizig_button = ttkb.Button(button_frame2, text="New Sizing", command=input,bootstyle="danger")
    new_sizig_button.grid(row=0, column=1, padx=5)

def quotation():
    quote_window = ttkb.Toplevel()
    quote_window.title("Quotation")
    global main_frame
    main_frame = ttkb.Frame(quote_window)
    main_frame.pack(expand=True, fill="both", padx=20, pady=20)
    global row_frames
    global text_areas
    row_frames = []
    text_areas = []
    def save_to_word():
        doc = docx.Document("Quote_format_High_Vtg.docx")
        if not doc.tables:
            tkmb.showerror("Error", "No table found in the template document.")
            return
        table = doc.tables[0]
        for row_idx, areas in enumerate(text_areas, start=1):
            while len(table.rows) <= row_idx:
                table.add_row()
            row_data = [area.get("1.0", tk.END).strip() for area in areas]
            if any(row_data):
                table.rows[row_idx].cells[0].text = str(row_idx)
            for col_idx, val in enumerate(row_data, start=1):
                if len(table.rows[row_idx].cells) > col_idx:
                    table.rows[row_idx].cells[col_idx].text = val
            # Calculate and set the product in the 6th column
            try:
                quantity = float(row_data[2]) if row_data[2] else 0
                price = float(row_data[3]) if row_data[3] else 0
                total = quantity * price
                if len(table.rows[row_idx].cells) > 5:
                    table.rows[row_idx].cells[5].text = str(round(total, 2))
            except Exception:
                if len(table.rows[row_idx].cells) > 5:
                    table.rows[row_idx].cells[5].text = ""
        save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
        if save_path:
            doc.save(save_path)
            tkmb.showinfo("Success", "Quotation saved to Word document.")

    # Create the button frame below main_frame
    button_frame = ttkb.Frame(quote_window)
    button_frame.pack(fill="x", padx=20, pady=(0, 20))
    save_button = ttkb.Button(button_frame, text="Save to Word", command=save_to_word)
    save_button.pack(padx=10, pady=10)

def option1():
    global option
    option=1
    add_row_frame()
def option2():
    global option
    option=2
    add_row_frame()
def option3():
    global option
    option=3
    add_row_frame()

def add_row_frame():
        row_frame = ttkb.Frame(main_frame)
        row_frame.pack(side="top", fill="x", pady=10)

        # Sr No label
        sr_no = len(row_frames) + 1
        sr_no_label = ttkb.Label(row_frame, text=f"{sr_no}", width=4, anchor="center")
        sr_no_label.pack(side="left", padx=(0, 5), pady=5)

        col1_frame = ttkb.Frame(row_frame)
        col1_frame.pack(side="left", expand=True, fill="both", padx=10)
        label1 = ttkb.Label(col1_frame, text="System")
        label1.pack(side="top", anchor="w")
        text_area1 = tk.Text(col1_frame, height=5, width=25, wrap="word")
        text_area1.pack(side="top", fill="both", expand=True)
        col2_frame = ttkb.Frame(row_frame)
        col2_frame.pack(side="left", expand=True, fill="both", padx=10)
        label2 = ttkb.Label(col2_frame, text="Solution")
        label2.pack(side="top", anchor="w")
        text_area2 = tk.Text(col2_frame, height=5, width=25, wrap="word")
        text_area2.pack(side="top", fill="both", expand=True)
        col3_frame = ttkb.Frame(row_frame)
        col3_frame.pack(side="left", expand=True, fill="both", padx=10)
        label3 = ttkb.Label(col3_frame, text="Quantity")
        label3.pack(side="top", anchor="w")
        text_area3 = tk.Text(col3_frame, height=5, width=25, wrap="word")
        text_area3.pack(side="top", fill="both", expand=True)
        col4_frame = ttkb.Frame(row_frame)
        col4_frame.pack(side="left", expand=True, fill="both", padx=10)
        label4 = ttkb.Label(col4_frame, text="Price")
        label4.pack(side="top", anchor="w")
        text_area4 = tk.Text(col4_frame, height=5, width=25, wrap="word")
        text_area4.pack(side="top", fill="both", expand=True)
        row_frames.append(row_frame)
        text_areas.append((text_area1, text_area2, text_area3, text_area4))
        text_area2.insert("1.0", f"Solution1: Lithium Battery Pack\n(HVL {offered_battery_config_value}) with\n Approximate Backup Time: {math.floor(backup_time)}Mins At BOL \n With Cabinet and inbuilt BMS")
        text_area3.insert("1.0", "1")
        if option == 1:
            if centretapping[0] == "centre tap":
                centretapping_text = "With Centre Tapping"
            else:
                centretapping_text = "Without Centre Tapping"
            text_area1.insert("1.0", f"{ups_rating}KVA: {backup_requirement}Min Backup \n(Load: {calc_load}kW)\n(Cell type:)\n({centretapping_text})")
            text_area4.insert("1.0", round(prices[0],2))
        elif option == 2:
            if centretapping[1] == "centre tap":
                centretapping_text = "With Centre Tapping"
            else:
                centretapping_text = "Without Centre Tapping"
            text_area1.insert("1.0", f"{ups_rating}KVA: {backup_requirement}Min Backup \n(Load: {calc_load}kW)\n(Cell type:)\n({centretapping_text})")
            text_area4.insert("1.0", round(prices[1],2))
        elif option == 3:
            if centretapping[2] == "centre tap":
                centretapping_text = "With Centre Tapping"
            else:
                centretapping_text = "Without Centre Tapping"
            text_area1.insert("1.0", f"{ups_rating}KVA: {backup_requirement}Min Backup \n(Load: {calc_load}kW)\n(Cell type:)\n({centretapping_text})")
            text_area4.insert("1.0", round(prices[2],2))

def new_costing():
    def set_dollar_rate():
        global dollar_rate
        dollar_rate = dollar_rate_entry.get()
        try:
            dollar_rate = float(dollar_rate)
            tkmb.showinfo("Success", "Dollar rate set successfully.")
            dollar_rate_src.destroy()
            costing_input()
        except ValueError:
            tkmb.showerror("Error", "Please enter a valid number for the dollar rate.")
    dollar_rate_src=ttkb.Toplevel()
    dollar_rate_src.title("Dollar Rate")
    dollar_rate_label = ttkb.Label(dollar_rate_src, text="Dollar Rate:")
    dollar_rate_label.grid(row=0, column=0, padx=10, pady=5, sticky="e")
    dollar_rate_entry = ttkb.Entry(dollar_rate_src, width=30)
    dollar_rate_entry.grid(row=0, column=1, padx=10, pady=5)
    dollar_rate_button = ttkb.Button(dollar_rate_src, text="Set Rate",command=set_dollar_rate)
    dollar_rate_button.grid(row=1, column=0, columnspan=2, pady=10)
    def costing_input():
        new_costing_screen = ttkb.Toplevel()
        new_costing_screen.title("New Costing")

        def landingcost(*args):
            try:
                inr1 = float(new_costing_inr1_entry.get() or 0)
            except ValueError:
                inr1 = 0
            try:
                inr2 = float(new_costing_inr2_entry.get() or 0)
            except ValueError:
                inr2 = 0
            try:
                other = float(new_costing_total_other_entry.get() or 0)
            except ValueError:
                other = 0
            total = inr1 + inr2 + other
            new_costing_landing_cost_entry.delete(0, tk.END)
            new_costing_landing_cost_entry.insert(0, str(round(total, 2)))
            try:
                landing_cost = float(new_costing_landing_cost_entry.get() or 0)
                overhead = round(landing_cost * 0.1, 2)
                warranty= round(landing_cost * 0.1, 2)
                costofpack = landing_cost + overhead + warranty
                new_costing_labour_entry.delete(0, tk.END)
                new_costing_labour_entry.insert(0, str(overhead))
                new_costing_warranty_entry.delete(0, tk.END)
                new_costing_warranty_entry.insert(0, str(warranty))
                new_costing_total_cost_entry.delete(0, tk.END)
                new_costing_total_cost_entry.insert(0, str(round(costofpack, 2)))
                margin10 = round(costofpack * 0.1, 2)
                new_costing_margin10_entry.delete(0, tk.END)
                new_costing_margin10_entry.insert(0, str(margin10))
                est_sales_cost_b = costofpack + margin10
                new_costing_est_sales_b_entry.delete(0, tk.END)
                new_costing_est_sales_b_entry.insert(0, str(round(est_sales_cost_b, 2)))
                margin15 = round(costofpack * 0.15, 2)
                new_costing_margin15_entry.delete(0, tk.END)
                new_costing_margin15_entry.insert(0, str(margin15))
                est_sales_cost_b5 = costofpack + margin15
                new_costing_est_sales_b5_entry.delete(0, tk.END)
                new_costing_est_sales_b5_entry.insert(0, str(round(est_sales_cost_b5, 2)))
            except ValueError:
                landing_cost = 0
            try:
                kw= float(new_costing_kw_entry.get() or 0)
                kwcost=(costofpack / kw)/dollar_rate
                kwprofit1= (est_sales_cost_b / kw)/dollar_rate
                kwprofit2= (est_sales_cost_b5 / kw)/dollar_rate
                new_costing_perkw_cost_entry.delete(0, tk.END)
                new_costing_perkw_cost_entry.insert(0, str(round(kwcost, 2)))
                new_costing_perkw_profit1_entry.delete(0, tk.END)
                new_costing_perkw_profit1_entry.insert(0, str(round(kwprofit1, 2)))
                new_costing_perkw_profit2_entry.delete(0, tk.END)
                new_costing_perkw_profit2_entry.insert(0, str(round(kwprofit2, 2)))
            except ValueError:
                kw = 0    
        

        # --- Group: Battery Details ---
        battery_details_frame = ttkb.LabelFrame(new_costing_screen, text="Battery Details", bootstyle="info")
        battery_details_frame.grid(row=0, column=0, columnspan=2, padx=10, pady=10, sticky="ew")

        new_costing_voltagelabel = ttkb.Label(battery_details_frame, text="Voltage:")
        new_costing_voltagelabel.grid(row=0, column=0, padx=10, pady=5, sticky="e")
        new_costing_voltageentry = ttkb.Entry(battery_details_frame, width=30)
        new_costing_voltageentry.grid(row=0, column=1, padx=10, pady=5)
        try:
            new_costing_voltageentry.insert(0, nominal_dc_voltage)
        except:
            pass

        new_costing_capacitylabel = ttkb.Label(battery_details_frame, text="Ampere Capacity:")
        new_costing_capacitylabel.grid(row=1, column=0, padx=10, pady=5, sticky="e")
        new_costing_capacityentry = ttkb.Entry(battery_details_frame, width=30)
        new_costing_capacityentry.grid(row=1, column=1, padx=10, pady=5)
        try: 
            new_costing_capacityentry.insert(0, value_str)
        except:
            pass

        new_costing_kw_label = ttkb.Label(battery_details_frame, text="Calculated kW:")
        new_costing_kw_label.grid(row=2, column=0, padx=10, pady=5, sticky="e")
        new_costing_kw_entry = ttkb.Entry(battery_details_frame, width=30)
        new_costing_kw_entry.grid(row=2, column=1, padx=10, pady=5)
        try:
            kw_calc = (float(new_costing_voltageentry.get()) * float(new_costing_capacityentry.get()) )/ 1000
            new_costing_kw_entry.insert(0, round(kw_calc, 2))
        except Exception:
            pass

        new_costing_cell_voltage_label = ttkb.Label(battery_details_frame, text="Cell Voltage:")
        new_costing_cell_voltage_label.grid(row=3, column=0, padx=10, pady=5, sticky="e")
        new_costing_cell_voltage_entry = ttkb.Entry(battery_details_frame, width=30)
        new_costing_cell_voltage_entry.grid(row=3, column=1, padx=10, pady=5)

        new_costing_cell_capacity_label = ttkb.Label(battery_details_frame, text="Cell Capacity:")
        new_costing_cell_capacity_label.grid(row=4, column=0, padx=10, pady=5, sticky="e")
        new_costing_cell_capacity_entry = ttkb.Entry(battery_details_frame, width=30)
        new_costing_cell_capacity_entry.grid(row=4, column=1, padx=10, pady=5)

        # Combination of cells in series
        new_costing_series_label = ttkb.Label(battery_details_frame, text="Combination of Cells in Series:")
        new_costing_series_label.grid(row=5, column=0, padx=10, pady=5, sticky="e")
        new_costing_series_entry = ttkb.Entry(battery_details_frame, width=30)
        new_costing_series_entry.grid(row=5, column=1, padx=10, pady=5)

        new_costing_parallel_label = ttkb.Label(battery_details_frame, text="Combination of Cells in Parallel:")
        new_costing_parallel_label.grid(row=6, column=0, padx=10, pady=5, sticky="e")
        new_costing_parallel_entry = ttkb.Entry(battery_details_frame, width=30)
        new_costing_parallel_entry.grid(row=6, column=1, padx=10, pady=5)
        
        def update_total_cells(*args):
            try:
                series = int(new_costing_series_entry.get())
                parallel = int(new_costing_parallel_entry.get())
                total = series * parallel
                new_costing_total_cells_entry.delete(0, tk.END)
                new_costing_total_cells_entry.insert(0, str(total))
            except ValueError:
                new_costing_total_cells_entry.delete(0, tk.END)

        new_costing_series_entry.bind("<KeyRelease>", update_total_cells)
        new_costing_parallel_entry.bind("<KeyRelease>", update_total_cells)

        # --- Group: Cell Costing Details (1) ---
        cell_costing_frame = ttkb.LabelFrame(new_costing_screen, text="Cell Costing Details (1)", bootstyle="info")
        cell_costing_frame.grid(row=0, column=2, columnspan=2, padx=10, pady=10, sticky="ew")

        # Total No Of Cells
        new_costing_total_cells_label = ttkb.Label(cell_costing_frame, text="Total No Of Cells:")
        new_costing_total_cells_label.grid(row=0, column=0, padx=10, pady=5, sticky="e")
        new_costing_total_cells_entry = ttkb.Entry(cell_costing_frame, width=30)
        new_costing_total_cells_entry.grid(row=0, column=1, padx=10, pady=5)

        # FOB Cost Of Cells
        new_costing_fob_cost_label = ttkb.Label(cell_costing_frame, text="FOB Cost Of Cells: $")
        new_costing_fob_cost_label.grid(row=1, column=0, padx=10, pady=5, sticky="e")
        new_costing_fob_cost_entry = ttkb.Entry(cell_costing_frame, width=30)
        new_costing_fob_cost_entry.grid(row=1, column=1, padx=10, pady=5)

        def update_total_fob_cost(*args):
            try:
                total_cells = float(new_costing_total_cells_entry.get())
                fob_cost = float(new_costing_fob_cost_entry.get())
                total_fob = total_cells * fob_cost
                new_costing_total_fob_entry.delete(0, tk.END)
                new_costing_total_fob_entry.insert(0, str(round(total_fob, 2)))
            except ValueError:
                new_costing_total_fob_entry.delete(0, tk.END)
            try:
                fob_val = float(new_costing_total_fob_entry.get())
                customs_val = round(fob_val * 0.075, 2)
                new_costing_customs_entry.delete(0, tk.END)
                new_costing_customs_entry.insert(0, str(customs_val))
            except ValueError:
                new_costing_customs_entry.delete(0, tk.END)
            try:
                fob_val = float(new_costing_total_fob_entry.get())
                customs_val = float(new_costing_customs_entry.get())
                landed_cost = fob_val + customs_val
                new_costing_landed_entry.delete(0, tk.END)
                new_costing_landed_entry.insert(0, str(round(landed_cost, 2)))
            except ValueError:
                new_costing_landed_entry.delete(0, tk.END)
            try:
                landed_cost = float(new_costing_landed_entry.get())
                inr_cost = round(landed_cost * dollar_rate, 2)
                new_costing_inr1_entry.delete(0, tk.END)
                new_costing_inr1_entry.insert(0, str(inr_cost))
            except ValueError:
                new_costing_inr1_entry.delete(0, tk.END)
            
            landingcost()
            

        new_costing_total_cells_entry.bind("<KeyRelease>", update_total_fob_cost)
        new_costing_fob_cost_entry.bind("<KeyRelease>", update_total_fob_cost)

        # Total FOB Cost of Cells
        new_costing_total_fob_label = ttkb.Label(cell_costing_frame, text="Total FOB Cost of Cells: $")
        new_costing_total_fob_label.grid(row=2, column=0, padx=10, pady=5, sticky="e")
        new_costing_total_fob_entry = ttkb.Entry(cell_costing_frame, width=30)
        new_costing_total_fob_entry.grid(row=2, column=1, padx=10, pady=5)

        # Clearing & Customs
        new_costing_customs_label = ttkb.Label(cell_costing_frame, text="Clearing & Customs: $")
        new_costing_customs_label.grid(row=3, column=0, padx=10, pady=5, sticky="e")
        new_costing_customs_entry = ttkb.Entry(cell_costing_frame, width=30)
        new_costing_customs_entry.grid(row=3, column=1, padx=10, pady=5)

        # Total Landed cost In India
        new_costing_landed_label = ttkb.Label(cell_costing_frame, text="Total Landed cost In India: $")
        new_costing_landed_label.grid(row=4, column=0, padx=10, pady=5, sticky="e")
        new_costing_landed_entry = ttkb.Entry(cell_costing_frame, width=30)
        new_costing_landed_entry.grid(row=4, column=1, padx=10, pady=5)

        # Cost In INR( Rs 87)-(1)
        new_costing_inr1_label = ttkb.Label(cell_costing_frame, text="Cost In INR -(1): ₹")
        new_costing_inr1_label.grid(row=5, column=0, padx=10, pady=5, sticky="e")
        new_costing_inr1_entry = ttkb.Entry(cell_costing_frame, width=30)
        new_costing_inr1_entry.grid(row=5, column=1, padx=10, pady=5)

        # --- Group: BMS/PCM & Landed Cost (2) ---
        bms_landed_frame = ttkb.LabelFrame(new_costing_screen, text="BMS/PCM (2)", bootstyle="info")
        bms_landed_frame.grid(row=7, column=0, columnspan=2, padx=10, pady=10, sticky="ew")

        # BMS/PCM
        new_costing_bms_label = ttkb.Label(bms_landed_frame, text="BMS/PCM: $")
        new_costing_bms_label.grid(row=0, column=0, padx=10, pady=5, sticky="e")
        new_costing_bms_entry = ttkb.Entry(bms_landed_frame, width=30)
        new_costing_bms_entry.grid(row=0, column=1, padx=10, pady=5)

        def landedcost2(*args):
            try:
                bms_val = float(new_costing_bms_entry.get())
                customs2_val = float(new_costing_customs2_entry.get())
                landed2_cost = bms_val + customs2_val
                new_costing_landed2_entry.delete(0, tk.END)
                new_costing_landed2_entry.insert(0, str(round(landed2_cost, 2)))
            except ValueError:
                new_costing_landed2_entry.delete(0, tk.END)
            try:
                landed2_cost = float(new_costing_landed2_entry.get())
                inr2_cost = round(landed2_cost * dollar_rate, 2)
                new_costing_inr2_entry.delete(0, tk.END)
                new_costing_inr2_entry.insert(0, str(inr2_cost))
            except ValueError:
                new_costing_inr2_entry.delete(0, tk.END)
            landingcost()

        def update_customs2(*args):
            try:
                bms_val = float(new_costing_bms_entry.get())
                customs2_val = round(bms_val * 0.2, 2)
                new_costing_customs2_entry.delete(0, tk.END)
                new_costing_customs2_entry.insert(0, str(customs2_val))
            except ValueError:
                new_costing_customs2_entry.delete(0, tk.END)
            landedcost2()
            landingcost()
            

        new_costing_bms_entry.bind("<KeyRelease>", update_customs2)
        

        # Clearing & Customs (again)
        new_costing_customs2_label = ttkb.Label(bms_landed_frame, text="Clearing & Customs: $")
        new_costing_customs2_label.grid(row=1, column=0, padx=10, pady=5, sticky="e")
        new_costing_customs2_entry = ttkb.Entry(bms_landed_frame, width=30)
        new_costing_customs2_entry.grid(row=1, column=1, padx=10, pady=5)
        new_costing_customs2_entry.bind("<KeyRelease>", landedcost2)

        # Total Landed cost In India (again)
        new_costing_landed2_label = ttkb.Label(bms_landed_frame, text="Total Landed cost In India: $")
        new_costing_landed2_label.grid(row=2, column=0, padx=10, pady=5, sticky="e")
        new_costing_landed2_entry = ttkb.Entry(bms_landed_frame, width=30)
        new_costing_landed2_entry.grid(row=2, column=1, padx=10, pady=5)

        # Cost In INR( Rs 87)-(2)
        new_costing_inr2_label = ttkb.Label(bms_landed_frame, text="Cost In INR -(2): ₹")
        new_costing_inr2_label.grid(row=3, column=0, padx=10, pady=5, sticky="e")
        new_costing_inr2_entry = ttkb.Entry(bms_landed_frame, width=30)
        new_costing_inr2_entry.grid(row=3, column=1, padx=10, pady=5)

        # --- Group: Other Components & Charges --- 
        other_components_frame = ttkb.LabelFrame(new_costing_screen, text="Other Components & Charges", bootstyle="info")
        other_components_frame.grid(row=7, column=2, columnspan=2, padx=10, pady=10, sticky="ew")

        # Cabinet (INR)
        new_costing_cabinet_label = ttkb.Label(other_components_frame, text="Cabinet (INR): ₹")
        new_costing_cabinet_label.grid(row=0, column=0, padx=10, pady=5, sticky="e")
        new_costing_cabinet_entry = ttkb.Entry(other_components_frame, width=30)
        new_costing_cabinet_entry.grid(row=0, column=1, padx=10, pady=5)

        # Bus Bar
        new_costing_busbar_label = ttkb.Label(other_components_frame, text="Bus Bar: ₹")
        new_costing_busbar_label.grid(row=1, column=0, padx=10, pady=5, sticky="e")
        new_costing_busbar_entry = ttkb.Entry(other_components_frame, width=30)
        new_costing_busbar_entry.grid(row=1, column=1, padx=10, pady=5)

        # Holder/caps
        new_costing_holder_label = ttkb.Label(other_components_frame, text="Holder/caps: ₹")
        new_costing_holder_label.grid(row=2, column=0, padx=10, pady=5, sticky="e")
        new_costing_holder_entry = ttkb.Entry(other_components_frame, width=30)
        new_costing_holder_entry.grid(row=2, column=1, padx=10, pady=5)

        # Wire & Gasket & Other Accessories
        new_costing_wire_gasket_label = ttkb.Label(other_components_frame, text="Wire & Gasket & Other Accessories: ₹")
        new_costing_wire_gasket_label.grid(row=3, column=0, padx=10, pady=5, sticky="e")
        new_costing_wire_gasket_entry = ttkb.Entry(other_components_frame, width=30)
        new_costing_wire_gasket_entry.grid(row=3, column=1, padx=10, pady=5)

        # Terminals+ Connectors
        new_costing_terminals_label = ttkb.Label(other_components_frame, text="Terminals+ Connectors: ₹")
        new_costing_terminals_label.grid(row=4, column=0, padx=10, pady=5, sticky="e")
        new_costing_terminals_entry = ttkb.Entry(other_components_frame, width=30)
        new_costing_terminals_entry.grid(row=4, column=1, padx=10, pady=5)

        # MCB/Fuse
        new_costing_mcb_label = ttkb.Label(other_components_frame, text="MCB/Fuse: ₹")
        new_costing_mcb_label.grid(row=5, column=0, padx=10, pady=5, sticky="e")
        new_costing_mcb_entry = ttkb.Entry(other_components_frame, width=30)
        new_costing_mcb_entry.grid(row=5, column=1, padx=10, pady=5)

        # Lugs & Slew
        new_costing_lugs_label = ttkb.Label(other_components_frame, text="Lugs & Slew: ₹")
        new_costing_lugs_label.grid(row=6, column=0, padx=10, pady=5, sticky="e")
        new_costing_lugs_entry = ttkb.Entry(other_components_frame, width=30)
        new_costing_lugs_entry.grid(row=6, column=1, padx=10, pady=5)

        # Nut Bolts
        new_costing_nutbolts_label = ttkb.Label(other_components_frame, text="Nut Bolts: ₹")
        new_costing_nutbolts_label.grid(row=7, column=0, padx=10, pady=5, sticky="e")
        new_costing_nutbolts_entry = ttkb.Entry(other_components_frame, width=30)
        new_costing_nutbolts_entry.grid(row=7, column=1, padx=10, pady=5)

        # Fiber glass + rod
        new_costing_fiberglass_label = ttkb.Label(other_components_frame, text="Fiber glass + rod: ₹")
        new_costing_fiberglass_label.grid(row=8, column=0, padx=10, pady=5, sticky="e")
        new_costing_fiberglass_entry = ttkb.Entry(other_components_frame, width=30)
        new_costing_fiberglass_entry.grid(row=8, column=1, padx=10, pady=5)

        # Awg cables
        new_costing_awg_label = ttkb.Label(other_components_frame, text="Awg cables: ₹")
        new_costing_awg_label.grid(row=9, column=0, padx=10, pady=5, sticky="e")
        new_costing_awg_entry = ttkb.Entry(other_components_frame, width=30)
        new_costing_awg_entry.grid(row=9, column=1, padx=10, pady=5)

        # Shipping Charges
        new_costing_shipping_label = ttkb.Label(other_components_frame, text="Shipping Charges: ₹")
        new_costing_shipping_label.grid(row=10, column=0, padx=10, pady=5, sticky="e")
        new_costing_shipping_entry = ttkb.Entry(other_components_frame, width=30)
        new_costing_shipping_entry.grid(row=10, column=1, padx=10, pady=5)

        # Packaging cost with safety packs
        new_costing_packaging_label = ttkb.Label(other_components_frame, text="Packaging cost with safety packs: ₹")
        new_costing_packaging_label.grid(row=11, column=0, padx=10, pady=5, sticky="e")
        new_costing_packaging_entry = ttkb.Entry(other_components_frame, width=30)
        new_costing_packaging_entry.grid(row=11, column=1, padx=10, pady=5)

        # --- Auto-calculate total for Other Chargers (3) ---
        def update_total_other(*args):
            try:
                cabinet = float(new_costing_cabinet_entry.get() or 0)
            except ValueError:
                cabinet = 0
            try:
                busbar = float(new_costing_busbar_entry.get() or 0)
            except ValueError:
                busbar = 0
            try:
                holder = float(new_costing_holder_entry.get() or 0)
            except ValueError:
                holder = 0
            try:
                wire_gasket = float(new_costing_wire_gasket_entry.get() or 0)
            except ValueError:
                wire_gasket = 0
            try:
                terminals = float(new_costing_terminals_entry.get() or 0)
            except ValueError:
                terminals = 0
            try:
                mcb = float(new_costing_mcb_entry.get() or 0)
            except ValueError:
                mcb = 0
            try:
                lugs = float(new_costing_lugs_entry.get() or 0)
            except ValueError:
                lugs = 0
            try:
                nutbolts = float(new_costing_nutbolts_entry.get() or 0)
            except ValueError:
                nutbolts = 0
            try:
                fiberglass = float(new_costing_fiberglass_entry.get() or 0)
            except ValueError:
                fiberglass = 0
            try:
                awg = float(new_costing_awg_entry.get() or 0)
            except ValueError:
                awg = 0
            try:
                shipping = float(new_costing_shipping_entry.get() or 0)
            except ValueError:
                shipping = 0
            try:
                packaging = float(new_costing_packaging_entry.get() or 0)
            except ValueError:
                packaging = 0

            total = (
            cabinet + busbar + holder + wire_gasket + terminals +
            mcb + lugs + nutbolts + fiberglass + awg + shipping + packaging
            )
            new_costing_total_other_entry.delete(0, tk.END)
            new_costing_total_other_entry.insert(0, str(round(total, 2)))
            
            
            
            landingcost()

        # Bind update to all entries
        for entry in [
            new_costing_cabinet_entry,
            new_costing_busbar_entry,
            new_costing_holder_entry,
            new_costing_wire_gasket_entry,
            new_costing_terminals_entry,
            new_costing_mcb_entry,
            new_costing_lugs_entry,
            new_costing_nutbolts_entry,
            new_costing_fiberglass_entry,
            new_costing_awg_entry,
            new_costing_shipping_entry,
            new_costing_packaging_entry
        ]:
            entry.bind("<KeyRelease>", update_total_other)

        # Total Other Chargers©(3)
        new_costing_total_other_label = ttkb.Label(other_components_frame, text="Total Other Chargers©(3): ₹")
        new_costing_total_other_label.grid(row=12, column=0, padx=10, pady=5, sticky="e")
        new_costing_total_other_entry = ttkb.Entry(other_components_frame, width=30)
        new_costing_total_other_entry.grid(row=12, column=1, padx=10, pady=5)
    
        new_costing_inr1_entry.bind("<KeyRelease>", landingcost)
        new_costing_inr2_entry.bind("<KeyRelease>", landingcost)
        new_costing_total_other_entry.bind("<KeyRelease>", landingcost)

        # --- Group: Cost Calculations & Margins ---
        cost_calc_frame = ttkb.LabelFrame(new_costing_screen, text="Cost Calculations & Margins", bootstyle="info")
        cost_calc_frame.grid(row=0, column=4, columnspan=2,rowspan=8, padx=10, pady=10, sticky="ew")

        # Landing cost of material (1+2+3)
        new_costing_landing_cost_label = ttkb.Label(cost_calc_frame, text="Landing cost of material (1+2+3): ₹")
        new_costing_landing_cost_label.grid(row=0, column=0, padx=10, pady=5, sticky="e")
        new_costing_landing_cost_entry = ttkb.Entry(cost_calc_frame, width=30)
        new_costing_landing_cost_entry.grid(row=0, column=1, padx=10, pady=5)

        # Production Labour & Assembly overheads
        new_costing_labour_label = ttkb.Label(cost_calc_frame, text="Production Labour & Assembly overheads: ₹")
        new_costing_labour_label.grid(row=1, column=0, padx=10, pady=5, sticky="e")
        new_costing_labour_entry = ttkb.Entry(cost_calc_frame, width=30)
        new_costing_labour_entry.grid(row=1, column=1, padx=10, pady=5)

        # Warranty & Service provision
        new_costing_warranty_label = ttkb.Label(cost_calc_frame, text="Warranty & Service provision: ₹")
        new_costing_warranty_label.grid(row=2, column=0, padx=10, pady=5, sticky="e")
        new_costing_warranty_entry = ttkb.Entry(cost_calc_frame, width=30)
        new_costing_warranty_entry.grid(row=2, column=1, padx=10, pady=5)

        # Total Cost of Pack (A)
        new_costing_total_cost_label = ttkb.Label(cost_calc_frame, text="Total Cost of Pack (A): ₹")
        new_costing_total_cost_label.grid(row=3, column=0, padx=10, pady=5, sticky="e")
        new_costing_total_cost_entry = ttkb.Entry(cost_calc_frame, width=30)
        new_costing_total_cost_entry.grid(row=3, column=1, padx=10, pady=5)

        # Margin @10 % On Cost
        new_costing_margin10_label = ttkb.Label(cost_calc_frame, text="Margin @10 % On Cost: ₹")
        new_costing_margin10_label.grid(row=4, column=0, padx=10, pady=5, sticky="e")
        new_costing_margin10_entry = ttkb.Entry(cost_calc_frame, width=30)
        new_costing_margin10_entry.grid(row=4, column=1, padx=10, pady=5)

        # Estimated Sales Cost-(B)
        new_costing_est_sales_b_label = ttkb.Label(cost_calc_frame, text="Estimated Sales Cost-(B): ₹")
        new_costing_est_sales_b_label.grid(row=5, column=0, padx=10, pady=5, sticky="e")
        new_costing_est_sales_b_entry = ttkb.Entry(cost_calc_frame, width=30)
        new_costing_est_sales_b_entry.grid(row=5, column=1, padx=10, pady=5)

        # Margin @15% On Cost
        new_costing_margin15_label = ttkb.Label(cost_calc_frame, text="Margin @15% On Cost: ₹")
        new_costing_margin15_label.grid(row=6, column=0, padx=10, pady=5, sticky="e")
        new_costing_margin15_entry = ttkb.Entry(cost_calc_frame, width=30)
        new_costing_margin15_entry.grid(row=6, column=1, padx=10, pady=5)

        # Estimated Sales Cost-(B+5)
        new_costing_est_sales_b5_label = ttkb.Label(cost_calc_frame, text="Estimated Sales Cost-(B+5): ₹")
        new_costing_est_sales_b5_label.grid(row=7, column=0, padx=10, pady=5, sticky="e")
        new_costing_est_sales_b5_entry = ttkb.Entry(cost_calc_frame, width=30)
        new_costing_est_sales_b5_entry.grid(row=7, column=1, padx=10, pady=5)

        # Per Kw Pricing @ cost (A)
        new_costing_perkw_cost_label = ttkb.Label(cost_calc_frame, text="Per Kw Pricing @ cost (A): $")
        new_costing_perkw_cost_label.grid(row=8, column=0, padx=10, pady=5, sticky="e")
        new_costing_perkw_cost_entry = ttkb.Entry(cost_calc_frame, width=30)
        new_costing_perkw_cost_entry.grid(row=8, column=1, padx=10, pady=5)

        # Per Kw pricing @ ist level profit (B)
        new_costing_perkw_profit1_label = ttkb.Label(cost_calc_frame, text="Per Kw pricing @ ist level profit (B): $")
        new_costing_perkw_profit1_label.grid(row=9, column=0, padx=10, pady=5, sticky="e")
        new_costing_perkw_profit1_entry = ttkb.Entry(cost_calc_frame, width=30)
        new_costing_perkw_profit1_entry.grid(row=9, column=1, padx=10, pady=5)

        # Per Kw pricing @ 2nd evel profit (B+5)
        new_costing_perkw_profit2_label = ttkb.Label(cost_calc_frame, text="Per Kw pricing @ 2nd evel profit (B+5): $")
        new_costing_perkw_profit2_label.grid(row=10, column=0, padx=10, pady=5, sticky="e")
        new_costing_perkw_profit2_entry = ttkb.Entry(cost_calc_frame, width=30)
        new_costing_perkw_profit2_entry.grid(row=10, column=1, padx=10, pady=5)

        def add_to_costing_table():
            # Collect all relevant entry values
            voltage = new_costing_voltageentry.get()
            capacity = new_costing_capacityentry.get()
            kw = new_costing_kw_entry.get()
            cell_voltage = new_costing_cell_voltage_entry.get()
            cell_capacity = new_costing_cell_capacity_entry.get()
            series = new_costing_series_entry.get()
            parallel = new_costing_parallel_entry.get()
            total_cells = new_costing_total_cells_entry.get()
            fob_cost = new_costing_fob_cost_entry.get()
            total_fob = new_costing_total_fob_entry.get()
            customs = new_costing_customs_entry.get()
            landed = new_costing_landed_entry.get()
            inr1 = new_costing_inr1_entry.get()
            bms = new_costing_bms_entry.get()
            customs2 = new_costing_customs2_entry.get()
            landed2 = new_costing_landed2_entry.get()
            inr2 = new_costing_inr2_entry.get()
            cabinet = new_costing_cabinet_entry.get()
            busbar = new_costing_busbar_entry.get()
            holder = new_costing_holder_entry.get()
            wire_gasket = new_costing_wire_gasket_entry.get()
            terminals = new_costing_terminals_entry.get()
            mcb = new_costing_mcb_entry.get()
            lugs = new_costing_lugs_entry.get()
            nutbolts = new_costing_nutbolts_entry.get()
            fiberglass = new_costing_fiberglass_entry.get()
            awg = new_costing_awg_entry.get()
            shipping = new_costing_shipping_entry.get()
            packaging = new_costing_packaging_entry.get()
            total_other = new_costing_total_other_entry.get()
            landing_cost = new_costing_landing_cost_entry.get()
            labour = new_costing_labour_entry.get()
            warranty = new_costing_warranty_entry.get()
            total_cost = new_costing_total_cost_entry.get()
            margin10 = new_costing_margin10_entry.get()
            est_sales_b = new_costing_est_sales_b_entry.get()
            margin15 = new_costing_margin15_entry.get()
            est_sales_b5 = new_costing_est_sales_b5_entry.get()
            perkw_cost = new_costing_perkw_cost_entry.get()
            perkw_profit1 = new_costing_perkw_profit1_entry.get()
            perkw_profit2 = new_costing_perkw_profit2_entry.get()

            # Prepare a summary string or tuple for the costing table
            
            costing_data = (
                voltage, capacity, kw, cell_voltage, cell_capacity, series, parallel, total_cells,
                fob_cost, total_fob, customs, landed, inr1, bms, customs2, landed2, inr2,
                cabinet, busbar, holder, wire_gasket, terminals, mcb, lugs, nutbolts, fiberglass,
                awg, shipping, packaging, total_other, landing_cost, labour, warranty, total_cost,
                margin10, est_sales_b, margin15, est_sales_b5, perkw_cost, perkw_profit1, perkw_profit2
            )
            # Find the first available Option column, or replace Option 1 if all are filled
            option_columns = ["Option 1", "Option 2", "Option 3"]
            # Prepare a flat list of costing_data as strings for display
            costing_data_str = [str(item) for item in costing_data]

            # Find the first empty option column in the first row, or replace Option 1 if all filled
            inserted = False
            for col in option_columns:
                # Check if the first row for this option is empty
                first_row_id = tree.get_children()[0]
                if not tree.set(first_row_id, col):
                    # Fill all rows for this option
                    for idx, item_id in enumerate(tree.get_children()):
                        if idx < len(costing_data_str):
                            tree.set(item_id, col, costing_data_str[idx])
                        else:
                            tree.set(item_id, col, "")
                    inserted = True
                    break
            if not inserted:
                # All options filled, replace Option 1
                for idx, item_id in enumerate(tree.get_children()):
                    if idx < len(costing_data_str):
                        tree.set(item_id, "Option 1", costing_data_str[idx])
                    else:
                        tree.set(item_id, "Option 1", "")

            tkmb.showinfo("Costing Table Entry", f"Added to costing table")

        add_to_costing_button = ttkb.Button(cost_calc_frame, text="Add To Costing Table", command=add_to_costing_table, bootstyle="success")
        add_to_costing_button.grid(row=9, column=4, columnspan=2, pady=10)



mainscreen()
root.mainloop()